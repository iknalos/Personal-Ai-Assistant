import asyncio
import shutil
from pathlib import Path
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

try:
    import pandas as pd
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_OK = True
except ImportError:
    WORD_OK = False

app = Server("windows-file-server")

# ── file type → folder name map ──────────────────────────────────────────────
TYPE_MAP = {
    "PDFs":        [".pdf"],
    "Word Docs":   [".doc", ".docx"],
    "Excel":       [".xls", ".xlsx", ".csv"],
    "PowerPoint":  [".ppt", ".pptx"],
    "Images":      [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".webp"],
    "Videos":      [".mp4", ".mov", ".avi", ".mkv", ".wmv"],
    "Audio":       [".mp3", ".wav", ".flac", ".aac"],
    "Archives":    [".zip", ".tar", ".gz", ".rar", ".7z"],
    "Code":        [".py", ".js", ".html", ".css", ".java", ".cpp", ".c", ".ts"],
    "Text":        [".txt", ".md", ".log"],
    "Executables": [".exe", ".msi", ".bat", ".sh"],
    "GIS Files":   [".gpkg", ".shp", ".geojson", ".kml"],
    "Notebooks":   [".ipynb"],
    "Data":        [".json", ".xml", ".yaml", ".yml"],
}
CHART_COLORS = ["#2196F3", "#4CAF50", "#FF5722", "#9C27B0", "#FF9800", "#00BCD4"]


def clean_df(df):
    """Remove fully-empty columns, unnamed columns, and fully-empty rows."""
    df = df.dropna(axis=1, how="all")
    df = df.loc[:, ~df.columns.astype(str).str.match(r"^Unnamed")]
    df = df.dropna(how="all").reset_index(drop=True)
    return df


def smart_read(path: Path, sheet=0) -> pd.DataFrame:
    """Read Excel/CSV and auto-detect header row and first column using openpyxl."""
    if path.suffix.lower() == ".csv":
        return clean_df(pd.read_csv(path))

    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[sheet] if isinstance(sheet, int) else wb[sheet]

    # Find first row AND first column that have real data
    header_row = None
    first_col   = None
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        non_empty = [(j, c) for j, c in enumerate(row) if c is not None and str(c).strip() != ""]
        if len(non_empty) >= 2:
            header_row = i          # 0-based row index for pandas header=
            first_col  = non_empty[0][0]  # 0-based column index
            break

    if header_row is None:
        raise ValueError("Could not find any data in the Excel file.")

    # Build list of column letters to use (skip empty leading columns)
    from openpyxl.utils import get_column_letter
    max_col = ws.max_column
    use_cols = [get_column_letter(first_col + 1) + ":" + get_column_letter(max_col)]

    df = pd.read_excel(path, sheet_name=sheet, header=header_row, usecols=use_cols[0])
    return clean_df(df)


@app.list_tools()
async def list_tools():
    return [
        Tool(name="list_files",
             description="List every file and folder inside a directory.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="read_file",
             description="Read the text contents of any file.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="create_folder",
             description="Create a new folder (including parent folders).",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="move_file",
             description="Move or rename a file.",
             inputSchema={"type": "object", "properties": {
                 "source": {"type": "string"}, "destination": {"type": "string"}
             }, "required": ["source", "destination"]}),

        Tool(name="copy_file",
             description="Copy a file to a new location.",
             inputSchema={"type": "object", "properties": {
                 "source": {"type": "string"}, "destination": {"type": "string"}
             }, "required": ["source", "destination"]}),

        Tool(name="delete_file",
             description="Delete a file or folder.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="search_files",
             description="Search recursively for files matching a pattern, e.g. *.xlsx",
             inputSchema={"type": "object", "properties": {
                 "path": {"type": "string"}, "pattern": {"type": "string"}
             }, "required": ["path", "pattern"]}),

        Tool(name="organize_folder",
             description="Organize all files in a folder into subfolders by file type.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="debug_excel",
             description="Debug an Excel file by showing raw cell contents row by row.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="read_excel",
             description="Read an Excel (.xlsx/.xls) or CSV file and return its data and statistics.",
             inputSchema={"type": "object", "properties": {
                 "path":  {"type": "string"},
                 "sheet": {"type": "string", "description": "Sheet name (optional)"}
             }, "required": ["path"]}),

        Tool(name="write_document",
             description="Create a Word .docx file with the given content and save it.",
             inputSchema={
                 "type": "object",
                 "properties": {
                     "path":    {"type": "string", "description": "Full path to save the .docx file"},
                     "title":   {"type": "string", "description": "Document title (optional)"},
                     "content": {"type": "string", "description": "Full document text content"},
                 },
                 "required": ["path", "content"]
             }),

        Tool(name="read_document",
             description="Read the text content of a Word .docx file.",
             inputSchema={"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]}),

        Tool(name="correct_document",
             description="Read a Word .docx file, apply corrected text, and save it.",
             inputSchema={
                 "type": "object",
                 "properties": {
                     "path":             {"type": "string", "description": "Path to the .docx file"},
                     "corrected_content":{"type": "string", "description": "The corrected full text"},
                 },
                 "required": ["path", "corrected_content"]
             }),

        Tool(name="generate_chart",
             description="Generate a chart (bar/line/pie/scatter) from an Excel or CSV file and save as PNG.",
             inputSchema={
                 "type": "object",
                 "properties": {
                     "path":        {"type": "string", "description": "Path to Excel/CSV file"},
                     "chart_type":  {"type": "string", "description": "bar | line | pie | scatter"},
                     "x_column":    {"type": "string", "description": "Column for X axis"},
                     "y_columns":   {"type": "array", "items": {"type": "string"}, "description": "Column(s) for Y axis"},
                     "title":       {"type": "string", "description": "Chart title (optional)"},
                     "output_path": {"type": "string", "description": "Where to save the PNG (optional)"}
                 },
                 "required": ["path", "chart_type", "x_column", "y_columns"]
             }),
    ]


@app.call_tool()
async def call_tool(name: str, arguments: dict):
    try:

        # ── list_files ────────────────────────────────────────────────────────
        if name == "list_files":
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"Path does not exist: {p}")
            items = sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower()))
            if not items:
                return _ok(f"Directory is empty: {p}")
            lines = []
            for item in items:
                tag  = "[FILE]" if item.is_file() else "[DIR] "
                size = f"  {item.stat().st_size:>12,} bytes" if item.is_file() else ""
                lines.append(f"{tag} {item.name}{size}")
            return _ok(f"Contents of {p}  ({len(items)} items)\n" + "\n".join(lines))

        # ── read_file ─────────────────────────────────────────────────────────
        elif name == "read_file":
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"File not found: {p}")
            return _ok(p.read_text(encoding="utf-8", errors="ignore"))

        # ── create_folder ─────────────────────────────────────────────────────
        elif name == "create_folder":
            p = Path(arguments["path"])
            p.mkdir(parents=True, exist_ok=True)
            return _ok(f"Folder created: {p}")

        # ── move_file ─────────────────────────────────────────────────────────
        elif name == "move_file":
            src, dst = Path(arguments["source"]), Path(arguments["destination"])
            if not src.exists():
                return _err(f"Source not found: {src}")
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src), str(dst))
            return _ok(f"Moved: {src} → {dst}")

        # ── copy_file ─────────────────────────────────────────────────────────
        elif name == "copy_file":
            src, dst = Path(arguments["source"]), Path(arguments["destination"])
            if not src.exists():
                return _err(f"Source not found: {src}")
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(src), str(dst))
            return _ok(f"Copied: {src} → {dst}")

        # ── delete_file ───────────────────────────────────────────────────────
        elif name == "delete_file":
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"Not found: {p}")
            shutil.rmtree(str(p)) if p.is_dir() else p.unlink()
            return _ok(f"Deleted: {p}")

        # ── search_files ──────────────────────────────────────────────────────
        elif name == "search_files":
            p = Path(arguments["path"])
            pattern = arguments["pattern"]
            if not p.exists():
                return _err(f"Path does not exist: {p}")
            matches = sorted(p.rglob(pattern))
            if not matches:
                return _ok(f"No files matching '{pattern}' found in {p}")
            return _ok(f"Found {len(matches)} file(s) matching '{pattern}':\n" + "\n".join(str(m) for m in matches))

        # ── organize_folder ───────────────────────────────────────────────────
        elif name == "organize_folder":
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"Path does not exist: {p}")
            moved, skipped = [], []
            for item in list(p.iterdir()):
                if not item.is_file():
                    continue
                ext = item.suffix.lower()
                folder = next((k for k, v in TYPE_MAP.items() if ext in v), "Other")
                target_dir = p / folder
                target_dir.mkdir(exist_ok=True)
                try:
                    shutil.move(str(item), str(target_dir / item.name))
                    moved.append(f"  {item.name}  →  {folder}/")
                except Exception as e:
                    skipped.append(f"  {item.name}  (skipped: {e})")
            if not moved and not skipped:
                return _ok("No files to organize.")
            out = f"Organized {len(moved)} file(s):\n" + "\n".join(moved)
            if skipped:
                out += f"\n\nSkipped {len(skipped)}:\n" + "\n".join(skipped)
            return _ok(out)

        # ── debug_excel ───────────────────────────────────────────────────────
        elif name == "debug_excel":
            import openpyxl
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"File not found: {p}")
            wb = openpyxl.load_workbook(p, data_only=True)
            ws = wb.active
            lines = [f"Sheet: {ws.title}  |  Dimensions: {ws.dimensions}"]
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 15:
                    lines.append("... (showing first 15 rows only)")
                    break
                lines.append(f"Row {i+1}: {list(row)}")
            return _ok("\n".join(lines))

        # ── read_excel ────────────────────────────────────────────────────────
        elif name == "read_excel":
            if not EXCEL_OK:
                return _err("Run: pip install pandas openpyxl matplotlib")
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"File not found: {p}")
            sheet = arguments.get("sheet", 0)
            df = smart_read(p, sheet)
            out  = f"File: {p}\n"
            out += f"Size: {df.shape[0]} rows × {df.shape[1]} columns\n"
            out += f"Columns: {', '.join(str(c) for c in df.columns)}\n\n"
            out += "First 10 rows:\n" + df.head(10).to_string(index=False)
            out += "\n\nSummary statistics:\n" + df.describe().to_string()
            return _ok(out)

        # ── write_document ────────────────────────────────────────────────────
        elif name == "write_document":
            if not WORD_OK:
                return _err("Run: pip install python-docx")
            p       = Path(arguments["path"])
            title   = arguments.get("title", "")
            content = arguments["content"]
            p.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            # Add title if provided
            if title:
                t = doc.add_heading(title, level=1)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add content — split by double newline for paragraphs
            for para in content.split('\n\n'):
                para = para.strip()
                if not para:
                    continue
                if para.startswith('# '):
                    doc.add_heading(para[2:], level=1)
                elif para.startswith('## '):
                    doc.add_heading(para[3:], level=2)
                elif para.startswith('### '):
                    doc.add_heading(para[4:], level=3)
                else:
                    doc.add_paragraph(para)
            doc.save(str(p))
            return _ok(f"Document saved to: {p}")

        # ── read_document ─────────────────────────────────────────────────────
        elif name == "read_document":
            if not WORD_OK:
                return _err("Run: pip install python-docx")
            p = Path(arguments["path"])
            if not p.exists():
                return _err(f"File not found: {p}")
            doc   = Document(str(p))
            lines = [para.text for para in doc.paragraphs if para.text.strip()]
            return _ok(f"Document: {p}\n\n" + "\n\n".join(lines))

        # ── correct_document ──────────────────────────────────────────────────
        elif name == "correct_document":
            if not WORD_OK:
                return _err("Run: pip install python-docx")
            p         = Path(arguments["path"])
            corrected = arguments["corrected_content"]
            if not p.exists():
                return _err(f"File not found: {p}")
            # Read existing doc to preserve title/heading style
            doc   = Document(str(p))
            title = ""
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    title = para.text
                    break
            # Rewrite with corrected content
            new_doc = Document()
            style   = new_doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            if title:
                h = new_doc.add_heading(title, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for para in corrected.split('\n\n'):
                para = para.strip()
                if para:
                    new_doc.add_paragraph(para)
            new_doc.save(str(p))
            return _ok(f"Document corrected and saved: {p}")

        # ── generate_chart ────────────────────────────────────────────────────
        elif name == "generate_chart":
            if not EXCEL_OK:
                return _err("Run: pip install pandas openpyxl matplotlib")
            p         = Path(arguments["path"])
            ctype     = arguments["chart_type"].lower()
            x_col     = arguments["x_column"]
            y_cols    = arguments["y_columns"]
            title     = arguments.get("title", f"{ctype.title()} Chart")
            out_path  = arguments.get("output_path") or str(p.parent / f"{p.stem}_{ctype}_chart.png")

            if not p.exists():
                return _err(f"File not found: {p}")

            df = smart_read(p)

            # Validate columns
            missing = [c for c in [x_col] + y_cols if c not in df.columns]
            if missing:
                return _err(
                    f"Column(s) not found: {missing}\n"
                    f"Available columns: {list(df.columns)}"
                )

            fig, ax = plt.subplots(figsize=(12, 7))
            fig.patch.set_facecolor("#FAFAFA")
            ax.set_facecolor("#F5F5F5")

            if ctype == "bar":
                w  = 0.8 / max(len(y_cols), 1)
                xs = range(len(df[x_col]))
                for i, col in enumerate(y_cols):
                    offset = (i - len(y_cols) / 2 + 0.5) * w
                    ax.bar([x + offset for x in xs], df[col], w,
                           label=col, color=CHART_COLORS[i % len(CHART_COLORS)], alpha=0.88)
                ax.set_xticks(list(xs))
                ax.set_xticklabels(df[x_col], rotation=45, ha="right")

            elif ctype == "line":
                for i, col in enumerate(y_cols):
                    ax.plot(df[x_col], df[col], marker="o", linewidth=2.5,
                            markersize=7, label=col, color=CHART_COLORS[i % len(CHART_COLORS)])
                plt.xticks(range(len(df[x_col])), df[x_col], rotation=45, ha="right")

            elif ctype == "pie":
                col = y_cols[0]
                ax.pie(df[col], labels=df[x_col], autopct="%1.1f%%",
                       startangle=90, colors=CHART_COLORS[:len(df)])

            elif ctype == "scatter":
                for i, col in enumerate(y_cols):
                    ax.scatter(df[x_col], df[col], label=col,
                               color=CHART_COLORS[i % len(CHART_COLORS)], s=90, alpha=0.75)

            ax.set_title(title, fontsize=16, fontweight="bold", pad=20)
            if ctype != "pie":
                ax.set_xlabel(x_col, fontsize=12)
                if len(y_cols) == 1:
                    ax.set_ylabel(y_cols[0], fontsize=12)
                ax.legend(fontsize=11)
                ax.grid(True, alpha=0.3)
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)

            plt.tight_layout()
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()
            return _ok(f"Chart saved to: {out_path}\n\nDouble-click the file in File Explorer to open it.")

        else:
            return _err(f"Unknown tool: {name}")

    except Exception as e:
        import traceback
        return _err(f"Error in {name}: {e}\n{traceback.format_exc()}")


# ── helpers ───────────────────────────────────────────────────────────────────
def _ok(text: str):  return [TextContent(type="text", text=text)]
def _err(text: str): return [TextContent(type="text", text=f"ERROR: {text}")]


async def main():
    print("=" * 50)
    print("  AI File & Excel Server ready")
    print("=" * 50)
    import sys
    # Filter out empty lines that cause JSON parse errors
    class FilteredStdin:
        def __init__(self, stream):
            self._stream = stream
        def __getattr__(self, name):
            return getattr(self._stream, name)
        async def receive(self, max_bytes=65536):
            while True:
                data = await self._stream.receive(max_bytes)
                if data and data.strip():
                    return data
    async with stdio_server() as (r, w):
        await app.run(r, w, app.create_initialization_options())

if __name__ == "__main__":
    asyncio.run(main())
